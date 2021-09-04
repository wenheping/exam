from django.db import models
from shiti.models import Shiti

from docxcompose.composer import Composer
from docx import Document

import fitz,uuid

file_path="/home/wen/exam/wenfiles/shiti/"

def w_uuid(string_length=10):
    random = str(uuid.uuid4())
    random = random.replace("-","")
    return random[0:string_length]

def w_bottom(page):
    # this function return the bottom position to add some text
    b_num=len(page.getTextBlocks())
    my_bottom=page.getTextBlocks()[0][3]
    if b_num>1 :
      for i in range(1,b_num,1):
        if page.getTextBlocks()[i][3]>my_bottom :
           my_bottom=page.getTextBlocks()[i][3]
    return my_bottom

def get_filename(zhangjie,test1,test2,test3):
    q=Shiti.objects.filter(w_zhishidian__contains=zhangjie)
    item_num=q.count()

    if item_num>0 :
      tmp_pdf=fitz.open()
      tmp_docx=Document(file_path+str(q[0].w_nian)+"/docx/"+q[0].w_timu+".docx")
      msg="("+str(q[0].w_nian)+"year"+q[0].w_shijuan_lei+")"
      para0=tmp_docx.paragraphs[0]
      para_tmp=para0.insert_paragraph_before(msg)
      composer=Composer(tmp_docx)

      for i in range(item_num):

         if test3=='true' and i>=1 :
           doc1=Document(file_path+str(q[i].w_nian)+"/docx/"+q[i].w_timu+".docx")
           msg="\n("+str(q[i].w_nian)+"year"+q[i].w_shijuan_lei+")"
           para0=doc1.paragraphs[0]
           para_tmp=para0.insert_paragraph_before(msg)
           composer.append(doc1)

         mypdf=fitz.open(file_path+str(q[i].w_nian)+"/pdf/"+q[i].w_timu+".pdf") # open origin pdf item
         tmp_pdf.insert_pdf(mypdf)
         r2=w_bottom(mypdf[0])

         pos1=fitz.Rect(80,50,400,500)
         pos2=fitz.Rect(105,50,400,500)
         tmp_pdf[i].insert_textbox(pos1,str(q[i].w_nian),color=(0,0,1))
         tmp_pdf[i].insert_textbox(pos2,"年"+q[i].w_shijuan_lei,fontname="china-ss",color=(0,0,1))
         # add the header to tell which year, which district

         if test1=='true':
           pos3=fitz.Rect(80,r2+30,500,842)
           tmp_pdf[i].insert_textbox(pos3,"答案：",fontname="china-ss",color=(1,0,0))
           pos3=fitz.Rect(90,r2+45,500,842)
           tmp_pdf[i].insert_textbox(pos3,q[i].w_daan,fontname="china-ss",fontsize=9,color=(1,0,0))
         # Add the key

         if test2=='true':
           r4=w_bottom(tmp_pdf[i])
           pos4=fitz.Rect(80,r4+20,500,842)
           tmp_pdf[i].insert_textbox(pos4,"解析：",fontname="china-ss",color=(1,0,0))
           pos4=fitz.Rect(80,r4+35,500,842)
           tmp_pdf[i].insert_textbox(pos4,q[i].w_jiexi,fontname="china-ss",fontsize=9,color=(1,0,0))
         # add the explation.

         mypdf.close()
      file_name=w_uuid()
      tmp_pdf.save(file_path+"tmp/"+file_name+".pdf")
      if test3=='true':
        composer.save(file_path+"tmp/"+file_name+".docx")

      return "tmp/"+file_name
    else:
      return "notfound"

def get_itemnum(zhangjie):
    q=Shiti.objects.filter(w_zhishidian__contains=zhangjie)
    num=q.count()
    return str(num)
